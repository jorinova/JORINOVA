"""
ALIS-X Document Reader Service
================================
Gives ALL AI services the ability to READ any document format.

Supported formats (offline-first, graceful degradation):
  ┌─────────────────┬────────────────────────────────────────────┐
  │ Format          │ Reader (in order of preference)            │
  ├─────────────────┼────────────────────────────────────────────┤
  │ PDF             │ pdfplumber → pypdf → Claude vision         │
  │ Word (.docx)    │ python-docx → zipfile XML extract          │
  │ Excel (.xlsx)   │ openpyxl → xlrd                            │
  │ CSV / TSV       │ stdlib csv (always available)              │
  │ Image (jpg/png) │ PIL + OCR → Claude vision                  │
  │ Audio           │ Whisper local                              │
  │ XML             │ stdlib xml.etree (always available)         │
  │ HL7 v2          │ hl7apy → regex parser                      │
  │ DICOM           │ pydicom                                    │
  │ JSON            │ stdlib json (always available)             │
  │ YAML            │ stdlib yaml (always available)             │
  │ TXT / LOG       │ stdlib open (always available)             │
  │ Markdown        │ regex strip → plain text                   │
  │ HTML            │ stdlib html.parser → plain text            │
  └─────────────────┴────────────────────────────────────────────┘

Design rules:
  - Every reader returns a ReadResult (never raises)
  - Offline readers are ALWAYS tried first
  - AI enhancement is OPTIONAL and non-blocking
  - Patient data is NEVER sent to cloud without explicit permission
  - Tables are extracted as list[list[str]] for structured parsing
"""
from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
import time
import xml.etree.ElementTree as ET
import yaml
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger('document_reader')


# ── Result schema ─────────────────────────────────────────────────────────────

@dataclass
class ReadResult:
    """Unified output for every document reader."""
    ok:          bool
    text:        str                   = ''       # full plain text
    pages:       list[str]            = field(default_factory=list)  # per-page text
    tables:      list[list[list[str]]]= field(default_factory=list)  # extracted tables
    metadata:    dict[str, Any]       = field(default_factory=dict)  # author, date, title, etc.
    structured:  dict[str, Any]       = field(default_factory=dict)  # key-value pairs extracted
    format:      str                  = ''       # detected format (pdf, docx, csv, …)
    reader_used: str                  = ''       # which reader succeeded
    ai_enhanced: bool                 = False    # True if AI ran extraction on top
    error:       str                  = ''
    warnings:    list[str]            = field(default_factory=list)
    latency_ms:  int                  = 0

    @property
    def word_count(self) -> int:
        return len(self.text.split()) if self.text else 0

    def to_dict(self) -> dict:
        d = self.__dict__.copy()
        return d


# ── Format detection ──────────────────────────────────────────────────────────

_EXT_MAP = {
    '.pdf':  'pdf',
    '.docx': 'docx',
    '.doc':  'doc',
    '.xlsx': 'xlsx',
    '.xls':  'xls',
    '.csv':  'csv',
    '.tsv':  'tsv',
    '.txt':  'txt',
    '.log':  'txt',
    '.xml':  'xml',
    '.hl7':  'hl7',
    '.json': 'json',
    '.yaml': 'yaml',
    '.yml':  'yaml',
    '.md':   'markdown',
    '.html': 'html',
    '.htm':  'html',
    '.jpg':  'image',
    '.jpeg': 'image',
    '.png':  'image',
    '.bmp':  'image',
    '.tiff': 'image',
    '.tif':  'image',
    '.webp': 'image',
    '.dcm':  'dicom',
    '.mp3':  'audio',
    '.wav':  'audio',
    '.webm': 'audio',
    '.ogg':  'audio',
    '.m4a':  'audio',
    '.flac': 'audio',
}

_MIME_MAP = {
    'application/pdf':                                     'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword':                                  'doc',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    'application/vnd.ms-excel':                            'xls',
    'text/csv':                                            'csv',
    'text/tab-separated-values':                           'tsv',
    'text/plain':                                          'txt',
    'text/xml': 'xml',
    'application/xml':                                     'xml',
    'text/html':                                           'html',
    'application/json':                                    'json',
    'application/dicom':                                   'dicom',
    'audio/mpeg':                                          'audio',
    'audio/wav':                                           'audio',
    'audio/webm':                                          'audio',
    'image/jpeg':                                          'image',
    'image/png':                                           'image',
    'image/tiff':                                          'image',
}


def detect_format(path: str, mime_type: str = '') -> str:
    """Detect document format from extension or MIME type."""
    if mime_type and mime_type in _MIME_MAP:
        return _MIME_MAP[mime_type]
    ext = Path(path).suffix.lower()
    if ext in _EXT_MAP:
        return _EXT_MAP[ext]
    # Peek at file header (magic bytes)
    try:
        with open(path, 'rb') as f:
            header = f.read(8)
        if header[:4] == b'%PDF':
            return 'pdf'
        if header[:2] == b'PK':            # ZIP-based: docx, xlsx
            # Check for OOXML signature
            with open(path, 'rb') as f:
                content = f.read(4096)
            if b'word/' in content:    return 'docx'
            if b'xl/'   in content:    return 'xlsx'
        if header[:4] == b'\xd0\xcf\x11\xe0':  # Compound Document (old .doc/.xls)
            return 'doc'
        if b'DICM' in header:
            return 'dicom'
    except Exception:
        pass
    return 'txt'


# ── Master entry point ────────────────────────────────────────────────────────

async def read(
    path:        str,
    mime_type:   str = '',
    fmt:         str = '',
    ai_enhance:  bool = False,    # run AI extraction on top of raw read
    cloud_ok:    bool = False,    # permit cloud AI for enhancement
    lang:        str = 'en',      # OCR language
    max_pages:   int = 50,        # limit for large PDFs
) -> ReadResult:
    """
    Universal document reader.
    Returns ReadResult — never raises.
    """
    t0 = time.time()
    path = str(path)

    if not Path(path).exists():
        return ReadResult(ok=False, error=f'File not found: {path}',
                          latency_ms=int((time.time()-t0)*1000))

    fmt = fmt or detect_format(path, mime_type)
    logger.info('Reading document: %s format=%s', Path(path).name, fmt)

    # Route to format-specific reader
    result = await _route_reader(path, fmt, lang, max_pages)
    result.format = fmt

    # AI enhancement (optional, non-blocking)
    if ai_enhance and result.ok and result.text:
        result = await _ai_enhance(result, cloud_ok, lang)

    result.latency_ms = int((time.time()-t0)*1000)
    return result


async def _route_reader(path: str, fmt: str, lang: str, max_pages: int) -> ReadResult:
    readers = {
        'pdf':      lambda: _read_pdf(path, max_pages),
        'docx':     lambda: _read_docx(path),
        'doc':      lambda: _read_doc_fallback(path),
        'xlsx':     lambda: _read_excel(path),
        'xls':      lambda: _read_excel(path),
        'csv':      lambda: _read_csv(path),
        'tsv':      lambda: _read_csv(path, delimiter='\t'),
        'txt':      lambda: _read_text(path),
        'xml':      lambda: _read_xml(path),
        'hl7':      lambda: _read_hl7(path),
        'json':     lambda: _read_json(path),
        'yaml':     lambda: _read_yaml(path),
        'markdown': lambda: _read_markdown(path),
        'html':     lambda: _read_html(path),
        'dicom':    lambda: _read_dicom(path),
        'image':    lambda: _read_image_text(path, lang),
        'audio':    lambda: _read_audio(path, lang),
    }
    reader = readers.get(fmt, lambda: _read_text(path))
    try:
        return await _run_sync(reader)
    except Exception as e:
        logger.error('Reader error for %s: %s', fmt, e)
        return ReadResult(ok=False, error=str(e), reader_used=fmt)


async def _run_sync(fn):
    """Run synchronous reader in executor to avoid blocking."""
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, fn)


# ── PDF reader ────────────────────────────────────────────────────────────────

def _read_pdf(path: str, max_pages: int = 50) -> ReadResult:
    # Try pdfplumber first (best table support)
    try:
        import pdfplumber
        pages_text, tables = [], []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages[:max_pages]):
                text = page.extract_text() or ''
                pages_text.append(text)
                # Extract tables
                for tbl in (page.extract_tables() or []):
                    cleaned = [[str(c or '').strip() for c in row] for row in tbl if any(c for c in row)]
                    if cleaned:
                        tables.append(cleaned)
        full_text = '\n\n'.join(pages_text)
        return ReadResult(ok=True, text=full_text, pages=pages_text,
                          tables=tables, reader_used='pdfplumber',
                          metadata={'page_count': len(pages_text)})
    except ImportError:
        logger.debug('pdfplumber not installed')
    except Exception as e:
        logger.warning('pdfplumber failed: %s', e)

    # Try pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        pages_text = []
        for page in reader.pages[:max_pages]:
            pages_text.append(page.extract_text() or '')
        meta = reader.metadata or {}
        return ReadResult(
            ok=True, text='\n\n'.join(pages_text), pages=pages_text,
            reader_used='pypdf',
            metadata={
                'page_count': len(reader.pages),
                'author':     str(meta.get('/Author', '')),
                'title':      str(meta.get('/Title', '')),
                'created':    str(meta.get('/CreationDate', '')),
            },
        )
    except ImportError:
        logger.debug('pypdf not installed')
    except Exception as e:
        logger.warning('pypdf failed: %s', e)

    # Fallback: raw binary read (last resort)
    try:
        raw = Path(path).read_bytes()
        # Naive text extraction: find readable ASCII runs
        text = re.sub(rb'[^\x20-\x7e\n\t]', b' ', raw).decode('ascii', errors='ignore')
        text = re.sub(r'\s+', ' ', text)[:5000]
        return ReadResult(ok=True, text=text, reader_used='raw_fallback',
                          warnings=['PDF parsed as raw bytes — install pdfplumber for full extraction'])
    except Exception as e:
        return ReadResult(ok=False, error=f'PDF read failed: {e}')


# ── Word DOCX reader ──────────────────────────────────────────────────────────

def _read_docx(path: str) -> ReadResult:
    try:
        import docx as python_docx
        doc       = python_docx.Document(path)
        paragraphs= [p.text for p in doc.paragraphs if p.text.strip()]
        tables    = []
        for tbl in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            tables.append(rows)
        props = doc.core_properties
        return ReadResult(
            ok=True, text='\n'.join(paragraphs), tables=tables,
            reader_used='python-docx',
            metadata={'author': props.author, 'title': props.title,
                      'created': str(props.created)},
        )
    except ImportError:
        logger.debug('python-docx not installed — trying ZIP fallback')
    except Exception as e:
        logger.warning('python-docx failed: %s', e)

    # ZIP XML fallback (docx is a ZIP file)
    # This is intentionally more robust than naive `tree.iter()` concatenation.
    # It extracts only `w:t` (text runs) and preserves paragraph breaks (w:p).
    try:
        import zipfile

        def _get_text_from_docx_xml(xml_root: ET.Element) -> str:
            # WordprocessingML uses namespaces; match by localname to avoid hardcoding.
            parts: list[str] = []

            def local(tag: str) -> str:
                return tag.split('}')[-1] if '}' in tag else tag

            # Iterate paragraphs in document order.
            for p in xml_root.iter():
                if local(p.tag) != 'p':
                    continue

                # Collect text runs under this paragraph.
                runs: list[str] = []
                for n in p.iter():
                    if local(n.tag) != 't':
                        continue
                    if n.text:
                        runs.append(n.text)

                paragraph_text = ''.join(runs)
                paragraph_text = paragraph_text.replace('\u00ad', '')
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text).strip()
                if paragraph_text:
                    parts.append(paragraph_text)

            # Fallback: if we couldn't detect paragraphs, at least collect all w:t runs.
            if not parts:
                all_runs: list[str] = []
                for n in xml_root.iter():
                    if local(n.tag) == 't' and n.text:
                        all_runs.append(n.text)
                text = ' '.join(all_runs)
                text = re.sub(r'\s+', ' ', text).strip()
                return text

            return '\n'.join(parts)

        with zipfile.ZipFile(path) as zf:
            if 'word/document.xml' in zf.namelist():
                xml_bytes = zf.read('word/document.xml')
                tree = ET.fromstring(xml_bytes)
                text = _get_text_from_docx_xml(tree)
                return ReadResult(
                    ok=True,
                    text=text,
                    reader_used='docx_zip_fallback',
                    warnings=['Install python-docx for full table extraction'],
                )
    except Exception as e:
        return ReadResult(ok=False, error=f'DOCX read failed: {e}')


    return ReadResult(ok=False, error='DOCX: no reader available')


def _read_doc_fallback(path: str) -> ReadResult:
    """Old .doc format — very limited without python-docx."""
    try:
        raw = Path(path).read_bytes()
        # Extract ASCII runs from binary .doc
        text = re.sub(rb'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', b' ', raw)
        text = text.decode('ascii', errors='ignore')
        text = re.sub(r'\s+', ' ', text)[:5000].strip()
        return ReadResult(ok=True, text=text, reader_used='doc_raw',
                          warnings=['Old .doc format — install python-docx for proper extraction'])
    except Exception as e:
        return ReadResult(ok=False, error=f'DOC read failed: {e}')


# ── Excel reader ──────────────────────────────────────────────────────────────

def _read_excel(path: str) -> ReadResult:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        all_tables, all_text = [], []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                clean = [str(c).strip() if c is not None else '' for c in row]
                rows.append(clean)
            if rows:
                all_tables.append(rows)
                all_text.append(f'=== Sheet: {sheet_name} ===')
                all_text.extend('\t'.join(r) for r in rows)
        return ReadResult(
            ok=True, text='\n'.join(all_text), tables=all_tables,
            reader_used='openpyxl',
            metadata={'sheets': wb.sheetnames},
        )
    except ImportError:
        logger.debug('openpyxl not installed')
    except Exception as e:
        logger.warning('openpyxl failed: %s', e)

    return ReadResult(ok=False, error='Excel: install openpyxl to read .xlsx/.xls files',
                      warnings=['pip install openpyxl'])


# ── CSV / TSV reader ──────────────────────────────────────────────────────────

def _read_csv(path: str, delimiter: str = ',') -> ReadResult:
    """Always available — uses Python stdlib csv."""
    try:
        rows = []
        encodings = ['utf-8', 'latin-1', 'cp1252']
        text = ''
        for enc in encodings:
            try:
                with open(path, newline='', encoding=enc) as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    rows   = list(reader)
                text = '\n'.join(delimiter.join(row) for row in rows)
                break
            except UnicodeDecodeError:
                continue
        if not rows:
            return ReadResult(ok=False, error='CSV: could not decode file')
        header   = rows[0] if rows else []
        data_rows= rows[1:] if len(rows) > 1 else []
        return ReadResult(
            ok=True, text=text, tables=[rows],
            reader_used='stdlib_csv',
            metadata={
                'row_count':  len(data_rows),
                'col_count':  len(header),
                'headers':    header,
                'delimiter':  delimiter,
            },
            structured={'headers': header, 'rows': data_rows[:100]},
        )
    except Exception as e:
        return ReadResult(ok=False, error=f'CSV read failed: {e}')


# ── Text reader ───────────────────────────────────────────────────────────────

def _read_text(path: str) -> ReadResult:
    """Always available — plain text with encoding detection."""
    for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
        try:
            text = Path(path).read_text(encoding=enc)
            return ReadResult(ok=True, text=text, reader_used='stdlib_text',
                              metadata={'encoding': enc, 'size_bytes': Path(path).stat().st_size})
        except UnicodeDecodeError:
            continue
    try:
        raw  = Path(path).read_bytes()
        text = raw.decode('ascii', errors='replace')
        return ReadResult(ok=True, text=text, reader_used='raw_ascii')
    except Exception as e:
        return ReadResult(ok=False, error=f'Text read failed: {e}')


# ── XML reader ────────────────────────────────────────────────────────────────

def _read_xml(path: str) -> ReadResult:
    """Always available — stdlib xml.etree."""
    try:
        tree = ET.parse(path)
        root = tree.getroot()

        def _flatten(node, depth=0) -> list[str]:
            parts = []
            tag   = node.tag.split('}')[-1] if '}' in node.tag else node.tag
            if node.text and node.text.strip():
                parts.append(f'{"  "*depth}{tag}: {node.text.strip()}')
            for child in node:
                parts.extend(_flatten(child, depth+1))
            return parts

        lines = _flatten(root)
        text  = '\n'.join(lines)

        # Extract as dict for structured access
        def _to_dict(node):
            d = {node.tag.split('}')[-1]: {}}
            d_inner = d[list(d.keys())[0]]
            if node.text and node.text.strip():
                d_inner['_text'] = node.text.strip()
            for k, v in node.attrib.items():
                d_inner[f'@{k}'] = v
            for child in node:
                key = child.tag.split('}')[-1]
                d_inner[key] = _to_dict(child)[key]
            return d

        return ReadResult(
            ok=True, text=text, reader_used='stdlib_xml',
            structured=_to_dict(root),
            metadata={'root_tag': root.tag.split('}')[-1], 'namespace': root.tag.split('}')[0].lstrip('{') if '}' in root.tag else ''},
        )
    except ET.ParseError as e:
        return ReadResult(ok=False, error=f'XML parse error: {e}')
    except Exception as e:
        return ReadResult(ok=False, error=f'XML read failed: {e}')


# ── HL7 v2 reader ─────────────────────────────────────────────────────────────

def _read_hl7(path: str) -> ReadResult:
    """Parse HL7 v2 messages. Uses hl7apy if available, else regex parser."""
    raw = _read_text(path)
    if not raw.ok:
        return raw
    text = raw.text

    # Try hl7apy
    try:
        from hl7apy.parser import parse_message
        msg = parse_message(text.replace('\n', '\r'))
        segments = {}
        for seg in msg.children:
            name = seg.name
            segments[name] = str(seg)
        return ReadResult(
            ok=True, text=text, reader_used='hl7apy',
            structured={'segments': segments, 'message_type': str(msg.MSH.MSH_9)},
            metadata={'hl7_version': str(msg.MSH.MSH_12)},
        )
    except ImportError:
        logger.debug('hl7apy not installed')
    except Exception as e:
        logger.debug('hl7apy parse error: %s', e)

    # Regex-based HL7 parser (offline fallback)
    try:
        segments = {}
        for line in re.split(r'[\r\n]+', text):
            line = line.strip()
            if not line:
                continue
            fields = line.split('|')
            seg_name = fields[0]
            segments[seg_name] = fields[1:]
        # Extract key fields
        msh = segments.get('MSH', [])
        pid = segments.get('PID', [])
        obr = segments.get('OBR', [])
        obx = segments.get('OBX', [])
        structured = {
            'message_type':   msh[8]  if len(msh) > 8  else '',
            'patient_name':   pid[4]  if len(pid) > 4  else '',
            'patient_dob':    pid[6]  if len(pid) > 6  else '',
            'test_name':      obr[4]  if len(obr) > 4  else '',
            'result_value':   obx[5]  if len(obx) > 5  else '',
            'result_units':   obx[6]  if len(obx) > 6  else '',
            'ref_range':      obx[7]  if len(obx) > 7  else '',
            'abnormal_flag':  obx[8]  if len(obx) > 8  else '',
            'all_segments':   segments,
        }
        return ReadResult(ok=True, text=text, reader_used='hl7_regex',
                          structured=structured)
    except Exception as e:
        return ReadResult(ok=False, error=f'HL7 parse failed: {e}')


# ── JSON reader ───────────────────────────────────────────────────────────────

def _read_json(path: str) -> ReadResult:
    try:
        raw = _read_text(path)
        if not raw.ok:
            return raw
        data = json.loads(raw.text)
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return ReadResult(ok=True, text=text, reader_used='stdlib_json',
                          structured=data if isinstance(data, dict) else {'data': data})
    except json.JSONDecodeError as e:
        return ReadResult(ok=False, error=f'JSON parse error: {e}')


# ── YAML reader ───────────────────────────────────────────────────────────────

def _read_yaml(path: str) -> ReadResult:
    try:
        raw  = _read_text(path)
        if not raw.ok:
            return raw
        data = yaml.safe_load(raw.text)
        text = yaml.dump(data, default_flow_style=False, allow_unicode=True)
        return ReadResult(ok=True, text=text, reader_used='stdlib_yaml',
                          structured=data if isinstance(data, dict) else {'data': data})
    except yaml.YAMLError as e:
        return ReadResult(ok=False, error=f'YAML parse error: {e}')


# ── Markdown reader ───────────────────────────────────────────────────────────

def _read_markdown(path: str) -> ReadResult:
    raw = _read_text(path)
    if not raw.ok:
        return raw
    # Strip markdown syntax to get plain text
    text = raw.text
    text = re.sub(r'#+\s*',     '',      text)  # headings
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*',     r'\1', text)  # italic
    text = re.sub(r'`{1,3}[^`]*`{1,3}', '', text)  # code
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # links
    text = re.sub(r'^[-*+]\s+', '', text, flags=re.MULTILINE)  # lists
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    raw.text = text
    raw.reader_used = 'markdown_strip'
    return raw


# ── HTML reader ───────────────────────────────────────────────────────────────

def _read_html(path: str) -> ReadResult:
    raw = _read_text(path)
    if not raw.ok:
        return raw
    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []
                self._skip = False
            def handle_starttag(self, tag, attrs):
                if tag in ('script', 'style'):
                    self._skip = True
            def handle_endtag(self, tag):
                if tag in ('script', 'style'):
                    self._skip = False
                if tag in ('p', 'div', 'br', 'h1', 'h2', 'h3', 'li', 'tr'):
                    self.parts.append('\n')
            def handle_data(self, data):
                if not self._skip and data.strip():
                    self.parts.append(data)

        parser = _TextExtractor()
        parser.feed(raw.text)
        text = ''.join(parser.parts)
        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        return ReadResult(ok=True, text=text, reader_used='stdlib_html')
    except Exception as e:
        return ReadResult(ok=False, error=f'HTML parse failed: {e}')


# ── DICOM reader ──────────────────────────────────────────────────────────────

def _read_dicom(path: str) -> ReadResult:
    try:
        import pydicom
        ds   = pydicom.dcmread(path, force=True)
        meta = {
            'patient_name':   str(getattr(ds, 'PatientName', '')),
            'patient_id':     str(getattr(ds, 'PatientID', '')),
            'patient_dob':    str(getattr(ds, 'PatientBirthDate', '')),
            'patient_sex':    str(getattr(ds, 'PatientSex', '')),
            'study_date':     str(getattr(ds, 'StudyDate', '')),
            'study_desc':     str(getattr(ds, 'StudyDescription', '')),
            'series_desc':    str(getattr(ds, 'SeriesDescription', '')),
            'modality':       str(getattr(ds, 'Modality', '')),
            'institution':    str(getattr(ds, 'InstitutionName', '')),
            'physician':      str(getattr(ds, 'ReferringPhysicianName', '')),
            'manufacturer':   str(getattr(ds, 'Manufacturer', '')),
        }
        # Get any text content (reports stored in DICOM SR)
        text_parts = [f'{k}: {v}' for k, v in meta.items() if v and v != 'None']
        text_content = getattr(ds, 'TextValue', '') or getattr(ds, 'ReportText', '')
        if text_content:
            text_parts.append(f'\nReport:\n{text_content}')
        return ReadResult(
            ok=True, text='\n'.join(text_parts), reader_used='pydicom',
            metadata=meta, structured=meta,
        )
    except ImportError:
        return ReadResult(ok=False, error='DICOM: install pydicom (pip install pydicom)',
                          warnings=['pip install pydicom'])
    except Exception as e:
        return ReadResult(ok=False, error=f'DICOM read failed: {e}')


# ── Image text reader (OCR) ───────────────────────────────────────────────────

def _read_image_text(path: str, lang: str = 'en') -> ReadResult:
    """Extract text from image using OCR. See ocr_service.py for full pipeline."""
    # Use Tesseract if available
    try:
        import pytesseract
        from PIL import Image
        from ai_services.language_service import get_tesseract_lang
        tess_lang = get_tesseract_lang(lang)
        img  = Image.open(path)
        text = pytesseract.image_to_string(img, lang=tess_lang)
        return ReadResult(ok=True, text=text.strip(), reader_used='tesseract',
                          metadata={'lang': tess_lang})
    except ImportError:
        logger.debug('pytesseract not installed')
    except Exception as e:
        logger.warning('Tesseract OCR error: %s', e)

    # EasyOCR fallback
    try:
        import easyocr
        from ai_services.language_service import get_whisper_code
        reader = easyocr.Reader([lang, 'en'], gpu=False, verbose=False)
        result = reader.readtext(path, detail=0)
        text   = ' '.join(result)
        return ReadResult(ok=True, text=text, reader_used='easyocr')
    except ImportError:
        logger.debug('easyocr not installed')
    except Exception as e:
        logger.warning('EasyOCR error: %s', e)

    # PIL description (at least returns image info)
    try:
        from PIL import Image
        img = Image.open(path)
        info = f'Image: {img.format} {img.size[0]}x{img.size[1]} {img.mode}'
        return ReadResult(
            ok=True, text=info, reader_used='pil_info',
            warnings=['No OCR engine installed. Install pytesseract or easyocr for text extraction.'],
            metadata={'format': img.format, 'width': img.size[0], 'height': img.size[1]},
        )
    except Exception as e:
        return ReadResult(ok=False, error=f'Image read failed: {e}')


# ── Audio reader (STT) ────────────────────────────────────────────────────────

def _read_audio(path: str, lang: str = 'en') -> ReadResult:
    """Transcribe audio using Whisper local model."""
    try:
        from ai_services.speech_service import transcribe_audio
        result = transcribe_audio(path, language=lang)
        if result.get('error'):
            return ReadResult(ok=False, error=result['error'])
        return ReadResult(
            ok=True, text=result.get('text', ''), reader_used='whisper',
            metadata={
                'language':    result.get('language', lang),
                'confidence':  result.get('confidence', 0.0),
                'latency_ms':  result.get('latency_ms', 0),
            },
        )
    except Exception as e:
        return ReadResult(ok=False, error=f'Audio read failed: {e}',
                          warnings=['Install openai-whisper for audio transcription'])


# ── AI enhancement layer ──────────────────────────────────────────────────────

async def _ai_enhance(result: ReadResult, cloud_ok: bool, lang: str) -> ReadResult:
    """
    Run AI extraction on top of raw text to extract structured data.
    Used for: lab reports, supplier invoices, clinical documents.
    """
    if not result.text or len(result.text) < 20:
        return result

    prompt = (
        f'Extract structured information from this {result.format} document.\n\n'
        f'Document text:\n{result.text[:3000]}\n\n'
        f'Extract: dates, names, values, test results, quantities, identifiers, totals.\n'
        f'Respond ONLY in JSON: {{"extracted":{{"key":"value"}},"summary":"one sentence"}}'
    )

    resp = None
    # Try local LLM first
    try:
        from ai_services.local_llm import generate, is_available
        if await is_available():
            resp = await generate(prompt, max_tokens=400, timeout_s=15.0)
    except Exception:
        pass

    # Try cloud if local failed and cloud is permitted
    if (not resp or not resp.content) and cloud_ok:
        try:
            from ai_services.cloud_llm import generate as cloud_gen, is_available as cloud_avail
            if await cloud_avail():
                resp = await cloud_gen(prompt, max_tokens=600)
        except Exception:
            pass

    if resp and resp.content:
        try:
            text = resp.content
            if '```json' in text:
                text = text.split('```json', 1)[1].rsplit('```', 1)[0]
            data = json.loads(text.strip())
            if 'extracted' in data:
                result.structured.update(data['extracted'])
            if 'summary' in data:
                result.metadata['ai_summary'] = data['summary']
            result.ai_enhanced = True
        except Exception as e:
            result.warnings.append(f'AI extraction parse error: {e}')

    return result


# ── Convenience helpers ───────────────────────────────────────────────────────

async def read_bytes(
    content:   bytes,
    filename:  str,
    mime_type: str = '',
    **kwargs,
) -> ReadResult:
    """Read from bytes — saves to temp file first."""
    import tempfile
    suffix = Path(filename).suffix or '.tmp'
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        f.write(content)
        tmp = f.name
    try:
        return await read(tmp, mime_type=mime_type, **kwargs)
    finally:
        Path(tmp).unlink(missing_ok=True)


def supported_formats() -> dict:
    """Return supported formats and their reader status."""
    checks = {
        'pdf':    ('pdfplumber', 'pypdf'),
        'docx':   ('docx',),
        'xlsx':   ('openpyxl',),
        'dicom':  ('pydicom',),
        'audio':  ('whisper',),
        'image_ocr': ('pytesseract', 'easyocr'),
        'hl7':    ('hl7apy',),
    }
    result = {}
    for fmt, pkgs in checks.items():
        installed = []
        for pkg in pkgs:
            try:
                __import__(pkg)
                installed.append(pkg)
            except ImportError:
                pass
        result[fmt] = {
            'available': bool(installed),
            'packages_installed': installed,
            'packages_required': list(pkgs),
            'fallback': fmt not in ('dicom', 'audio') or bool(installed),
        }
    # Always-available formats
    for fmt in ('csv', 'txt', 'json', 'yaml', 'xml', 'markdown', 'html', 'tsv'):
        result[fmt] = {'available': True, 'packages_installed': ['stdlib'], 'fallback': True}
    return result
